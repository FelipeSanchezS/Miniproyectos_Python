from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

class WordAutomation:
    def __init__(self, filename="documento_generado.docx"):
        """Inicializa un documento de Word."""
        self.filename = filename
        self.document = Document()

    def agregar_párrafo(self, texto, fuente = "Arial", tamaño = 12):
        """Agrega un párrafo en especifico"""
        p=self.document.add_paragraph(texto)
        run = p.runs[0]
        run.font.name = fuente
        run.font.size = Pt(tamaño)

    def agregar_tabla_desde_excel(self, excel_file):
        """Importa los datos de un archivo excel y los inserta en una tabla del documento"""
        df = pd.read_excel(excel_file)

        table = self.document.add_table(rows=1, cols=len(df.columns))
        hdr_cells = table.rows[0].cells

        #Asignar los nombres de las columnas de la tabla
        for i, column in enumerate(df.columns):
            hdr_cells[i].text = column

        #Asignamos a cada fila los valores del dataframe
        for _, row in df.iterrows():
             row_cells = table.add_row().cells
             for i, item in enumerate(row):
                 row_cells[i].text = str(item)
    
    def guardar_documento(self):
        """Guarda el documento en el archivo especificado"""
        self.document.save(self.filename)
        print(f"Documento guardado como: {self.filename}")
