from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

#Creamos el documento
document = Document()

#Primer Párrafo - va estar centrado con diferentes estilos en el mismo párrafo
p1 = document.add_paragraph()

#Con esto centramos el párrafo
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

#Agregamos primer texto al párrafo
run1 = p1.add_run("Este es un texto")

#definimos otros estilos
run1.font.name = 'Arial' #Definimos letra
run1.font.size = Pt(12) #Definimos tamaño

#Guardamos el documento
document.save("AutomatizaciónV2.docx")