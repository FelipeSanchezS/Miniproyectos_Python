from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

#Creamos el documento
document = Document()



#------PÁRRAFO 1-------
#En esta parte creamos los diferentes párrafos que vayamos a necesitar
párrafo1 = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo1.alignment = WD_ALIGN_PARAGRAPH.CENTER
#Agregamos primer texto al párrafo y definimos otros estilos
run = párrafo1.add_run("Este es un texto.") 
run.font.name = 'Arial' #Definimos letra
run.font.size = Pt(12) #Definimos tamaño



#------PÁRRAFO 2-------
#En esta parte creamos los diferentes párrafos que vayamos a necesitar
párrafo2 = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo2.alignment = WD_ALIGN_PARAGRAPH.LEFT
#Agregamos Segundo párrafo y definimos otros estilos
run = párrafo2.add_run("Con otro tipo de letra se diseño la segunda linea hacia la izquierda - ") 
run.font.name = 'Times New Roman' #Definimos letra
run.font.size = Pt(14) #Definimos tamaño
run.font.bold = True
run = párrafo2.add_run("Es la continuación de la misma segunda linea con otros estilos para seguir probando.") 
run.font.name = 'Calibri' #Definimos letra
run.font.size = Pt(18) #Definimos tamaño
run.font.bold = True



#------PÁRRAFO 3-------
#En esta parte creamos los diferentes párrafos que vayamos a necesitar
párrafo3 = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
#Agregamos párrafo 3 y agregamos textos
run = párrafo3.add_run("Este es el estilo correspondiente al tercer párrafo del documento que se esta creando.") 
run.font.name = 'Georgia' #Definimos letra
run.font.size = Pt(11) #Definimos tamaño
run.font.color.rgb = RGBColor(0,0,255)



#------PÁRRAFO 4-------
#En esta parte creamos los diferentes párrafos que vayamos a necesitar
párrafo4 = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
#Agregamos párrafo 4 y agregamos textos
run = párrafo4.add_run("Este es el siguiente párrafo, (párrafo 4) creado con los diferentes estilos ya predefinidos") 
run.font.name = 'Verdana' #Definimos letra
run.font.size = Pt(13) #Definimos tamaño
run.font.color.rgb = RGBColor(0,0,255)



#------PÁRRAFO 5-------
nombre = input('Ingrese su nombre: ')
#En esta parte creamos los diferentes párrafos que vayamos a necesitar
párrafo5 = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo5.alignment = WD_ALIGN_PARAGRAPH.LEFT
#Agregamos párrafo 4 y agregamos textos
run = párrafo5.add_run("Este es el siguiente párrafo, (párrafo 5) creado con los diferentes estilos ya predefinidos. por ultimo, su nombre es "+nombre) 
run.font.name = 'Courier New' #Definimos letra
run.font.size = Pt(13) #Definimos tamaño
run.font.italic = True




#Guardamos el documento
document.save("AutomatizaciónV2.docx")