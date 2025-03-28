#Importamos librerías
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import io
import openpyxl

#Función para crear el reporte o informe
def create_report(template_path, data, chart_data=None):
    st.write("Iniciando la creación del informe")
    #cargamos documento word
    doc = Document(template_path)
    #vamos a validar cara párrafo del documento
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            #validamos las llaves {{}} del arreglo e ingresamos los valores que sea necesarios cuando identifican valores
            if f'{{{{{key}}}}}' in paragraph.text:
                st.write(f"Reemplazando {key} con {value} en el informe.")
            paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', str(value))

    #Sección de crear el gráfico
    if chart_data is not None:
        st.write("Generando el gráfico")
        plt.figure(figsize=(6,4))
        plt.bar(chart_data['labels'], chart_data['values'])
        plt.title(chart_data['title'])
        plt.xlabel(chart_data.get('xlabel', ''))
        plt.ylabel(chart_data.get('ylabel', ''))
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format = 'png')
        img_buffer.seek(0)
        st.write("Insertando gráfico en el marcador del documento!")

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if '[Aquí se insertará el grafico]' in run.text:
                    run.text = run.text.replace('[Aquí se insertará el grafico]', '')
                    run.add_picture(img_buffer, width=Inches(6))

    
    #Guardamos en la memoria
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    st.write("Reporte creado con éxito!")

    return output




def main():
    #titulo
    st.title("Generador de informes mediante python")
    #Acá se pone para cargar la plantilla word
    template_file = st.file_uploader("Cargue su plantilla Word", type="docx")

    #Acá se pone para cargar la el archivo de excel ocn la información
    data_file = st.file_uploader("Cargue sus datos", type=["xlsx", "cvs"])

    if template_file and data_file:
        st.success("Archivos cargados de forma correcta")

        #Cargamos el archivo con los datos que se subieron desde el pc
        # se usa el read cvs si es .cvs, pero si es excel se usa el read_excel
        df = pd.read_csv(data_file) if data_file.name.endswith('.csv') else pd.read_excel(data_file) 

        #creamos encabezado
        st.subheader("Datos cargados")
        #Mostramos los datos que están en el dataset
        st.dataframe(df)

        #Seleccionamos la fila que queremos que este en el informe
        row_index = st.selectbox("Seleccione la fila para crear el informe: ", options=range(len(df)))
        #guardamos la información escogida dentro de una variable
        selected_data = df.iloc[row_index].to_dict()
        # st.write(selected_data)

        #Sección para agregar el gráfico
        generate_chart = st.checkbox("Generar grafico: ")
        chart_data = None

        if generate_chart:
            chart_title = st.text_input("Titulo del gráfico" , "Gráfico de datos" )
            x_column = st.selectbox("Columna para eje X", options=df.columns)
            y_column = st.selectbox("Columna para eje Y", options=df.columns)

            #Aca entregamos los datos del dataset a la gráfica
            chart_data = {
                'title': chart_title,
                'labels': df[x_column].tolist(),
                'values': df[y_column].tolist(),
                'xlabel': x_column,
                'ylabel':y_column
            }

            st.write("Los datos del gráfico son: ", chart_data)

        #Botón para generar informe, al darle click me lleva la función de create_report
        if st.button("Generar informe"):
            output = create_report(template_file, selected_data, chart_data)
            #Creamos botón de descarga
            st.download_button("Descarga tu informe ", output, "informe-generado.docx","application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            

main()