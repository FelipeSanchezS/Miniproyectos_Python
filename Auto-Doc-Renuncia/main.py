#Importamos librerias
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

#Creamos el documento
document = Document()

#Ingresamos los datos necesarios
nombre = input('Ingrese su nombre: ')
dirección = input('Ingrese su dirección de residencia: ')
email = input('Ingrese su correo: ')
teléfono = input('Ingrese su teléfono: ')
fechaA = input('Ingrese la fecha: ')
empresa = input('Ingrese nombre de la empresa: ')
ciudadPaís = input('Ingrese ciudad y país: ')
cargo = input("Ingrese su cargo desempeñado: ")
fechaF = input("Ingrese la fecha de efectividad de la renuncia: ")

#------Primera linea-------
párrafo = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
#Agregamos párrafo 4 y agregamos textos
run = párrafo.add_run(nombre) 
run.font.name = 'Arial' #Definimos letra
run.font.size = Pt(11)
run.font.italic = True

#------Segunda linea-------
párrafo = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
#Agregamos párrafo 4 y agregamos textos
run = párrafo.add_run(dirección) 
run.font.name = 'Arial' #Definimos letra
run.font.size = Pt(11)
run.font.italic = True

#------Tercera linea-------
párrafo = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
#Agregamos párrafo 4 y agregamos textos
run = párrafo.add_run(email) 
run.font.name = 'Arial' #Definimos letra
run.font.size = Pt(11)
run.font.italic = True

#------Cuarta linea-------
párrafo = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
#Agregamos párrafo 4 y agregamos textos
run = párrafo.add_run(teléfono) 
run.font.name = 'Arial' #Definimos letra
run.font.size = Pt(11)
run.font.italic = True

#------Quinta linea-------
párrafo = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
#Agregamos párrafo 4 y agregamos textos
run = párrafo.add_run(fechaA) 
run.font.name = 'Arial' #Definimos letra
run.font.size = Pt(11)
run.font.italic = True

#------Salto de linea-------
document.add_paragraph()

#------Sexta linea-------
párrafo = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
#Agregamos párrafo 4 y agregamos textos
run = párrafo.add_run(empresa) 
run.font.name = 'Arial' #Definimos letra
run.font.size = Pt(11)
run.font.italic = True

#------Séptima linea-------
párrafo = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
#Agregamos párrafo 4 y agregamos textos
run = párrafo.add_run(ciudadPaís) 
run.font.name = 'Arial' #Definimos letra
run.font.size = Pt(11)
run.font.italic = True

#------Salto de linea-------
document.add_paragraph()

#------Octava linea-------
párrafo = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
#Agregamos párrafo 4 y agregamos textos
run = párrafo.add_run("Estimado/a "+empresa+": ") 
run.font.name = 'Arial' #Definimos letra
run.font.size = Pt(11)

#------Primer párrafo de la carta-------
párrafo = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
#Agregamos párrafo 4 y agregamos textos
run = párrafo.add_run("Por medio de la presente, me dirijo a usted para presentar mi renuncia al cargo "+ cargo +" en "+ empresa +", con efectividad a partir del "+ fechaF) 
run.font.name = 'Arial' #Definimos letra
run.font.size = Pt(11) #Definimos tamaño

#------Salto de linea-------
document.add_paragraph()

#------Segundo párrafo de la carta-------
párrafo = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
#Agregamos párrafo 4 y agregamos textos
run = párrafo.add_run("""Ha sido un placer formar parte de esta empresa y agradezco todas las oportunidades que se me han brindado durante mi tiempo aquí. Estoy agradecido/a por el apoyo y la experiencia que he adquirido, los cuales han contribuido significativamente a mi crecimiento profesional.""") 
run.font.name = 'Arial' #Definimos letra
run.font.size = Pt(11) #Definimos tamaño

#------Salto de linea-------
document.add_paragraph()

#------Tercer párrafo de la carta-------
párrafo = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
#Agregamos párrafo 4 y agregamos textos
run = párrafo.add_run("""Durante mi período, me comprometo a asegurar una transición fluida en mis responsabilidades y a colaborar en cualquier proceso que facilite la transferencia de mis tareas a mi sucesor/a. """) 
run.font.name = 'Arial' #Definimos letra
run.font.size = Pt(11) #Definimos tamaño

#------Salto de linea-------
document.add_paragraph()

#------Cuarto párrafo de la carta-------
párrafo = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
#Agregamos párrafo 4 y agregamos textos
run = párrafo.add_run("""Agradezco su comprensión y espero poder mantener el contacto en el futuro. """) 
run.font.name = 'Arial' #Definimos letra
run.font.size = Pt(11) #Definimos tamaño

#------Salto de linea-------
document.add_paragraph()

#------Quinto párrafo de la carta-------
párrafo = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
#Agregamos párrafo 4 y agregamos textos
run = párrafo.add_run("""Atentamente, """) 
run.font.name = 'Arial' #Definimos letra
run.font.size = Pt(11) #Definimos tamaño

#------Salto de linea-------
document.add_paragraph()

#------Quinto párrafo de la carta-------
párrafo = document.add_paragraph()
#Con esto elegimos dirección del párrafo el párrafo
párrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
#Agregamos párrafo 4 y agregamos textos
run = párrafo.add_run(nombre) 
run.font.name = 'Arial' #Definimos letra
run.font.size = Pt(11) #Definimos tamaño

#Guardamos el documento
document.save("Carta_Renuncia "+nombre+".docx")
print('Documento creado de forma correcta')